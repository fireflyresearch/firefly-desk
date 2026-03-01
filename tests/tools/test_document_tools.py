# Copyright 2026 Firefly Software Solutions Inc
#
# Licensed under the Apache License, Version 2.0

"""Tests for document tools (read, create, modify, convert)."""

from __future__ import annotations

import io
from unittest.mock import AsyncMock, MagicMock

import pytest

from flydesk.catalog.enums import RiskLevel
from flydesk.tools.builtin import BUILTIN_SYSTEM_ID, BuiltinToolRegistry
from flydesk.tools.document_tools import (
    DocumentToolExecutor,
    document_convert_tool,
    document_create_tool,
    document_modify_tool,
    document_read_tool,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def storage() -> AsyncMock:
    """Mock FileStorageProvider."""
    mock = AsyncMock()
    mock.store = AsyncMock(return_value="/files/output.bin")
    mock.retrieve = AsyncMock(return_value=b"")
    return mock


@pytest.fixture
def executor(storage: AsyncMock) -> DocumentToolExecutor:
    return DocumentToolExecutor(storage=storage)


def _make_docx_bytes(paragraphs: list[str] | None = None) -> bytes:
    """Create a minimal .docx file in memory."""
    from docx import Document

    doc = Document()
    for para in paragraphs or ["Hello world"]:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_xlsx_bytes(
    rows: list[list] | None = None, sheet_name: str = "Sheet1"
) -> bytes:
    """Create a minimal .xlsx file in memory."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    for row in rows or [["Name", "Age"], ["Alice", 30]]:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_pdf_bytes(text: str = "Hello PDF") -> bytes:
    """Create a minimal PDF in memory."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    doc.build([Paragraph(text, styles["BodyText"])])
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Tool definition tests
# ---------------------------------------------------------------------------


class TestToolDefinitions:
    def test_document_read_tool_definition(self):
        tool = document_read_tool()
        assert tool.name == "document_read"
        assert tool.risk_level == RiskLevel.READ
        assert tool.system_id == BUILTIN_SYSTEM_ID
        assert "file_path" in tool.parameters

    def test_document_create_tool_definition(self):
        tool = document_create_tool()
        assert tool.name == "document_create"
        assert tool.risk_level == RiskLevel.LOW_WRITE
        assert "format" in tool.parameters
        assert "title" in tool.parameters

    def test_document_modify_tool_definition(self):
        tool = document_modify_tool()
        assert tool.name == "document_modify"
        assert tool.risk_level == RiskLevel.LOW_WRITE
        assert "operation" in tool.parameters

    def test_document_convert_tool_definition(self):
        tool = document_convert_tool()
        assert tool.name == "document_convert"
        assert tool.risk_level == RiskLevel.READ
        assert "target_format" in tool.parameters


class TestRegistryIntegration:
    def test_admin_gets_document_tools(self):
        """Admin user with * permission gets all document tools."""
        tools = BuiltinToolRegistry.get_tool_definitions(["*"])
        names = {t.name for t in tools}
        assert "document_read" in names
        assert "document_create" in names
        assert "document_modify" in names
        assert "document_convert" in names

    def test_knowledge_read_gets_read_tools_only(self):
        """knowledge:read permission grants read/convert but not create/modify."""
        tools = BuiltinToolRegistry.get_tool_definitions(["knowledge:read"])
        names = {t.name for t in tools}
        assert "document_read" in names
        assert "document_convert" in names
        assert "document_create" not in names
        assert "document_modify" not in names

    def test_knowledge_write_gets_write_tools(self):
        """knowledge:write permission grants create/modify tools."""
        tools = BuiltinToolRegistry.get_tool_definitions(["knowledge:write"])
        names = {t.name for t in tools}
        assert "document_create" in names
        assert "document_modify" in names

    def test_no_permission_excludes_document_tools(self):
        """User with no special permissions does not get document tools."""
        tools = BuiltinToolRegistry.get_tool_definitions([])
        names = {t.name for t in tools}
        assert "document_read" not in names
        assert "document_create" not in names


class TestIsDocumentTool:
    def test_recognized_tools(self):
        assert DocumentToolExecutor.is_document_tool("document_read") is True
        assert DocumentToolExecutor.is_document_tool("document_create") is True
        assert DocumentToolExecutor.is_document_tool("document_modify") is True
        assert DocumentToolExecutor.is_document_tool("document_convert") is True

    def test_unrecognized_tools(self):
        assert DocumentToolExecutor.is_document_tool("search_knowledge") is False
        assert DocumentToolExecutor.is_document_tool("document_unknown") is False


# ---------------------------------------------------------------------------
# Read tests
# ---------------------------------------------------------------------------


class TestDocumentRead:
    async def test_read_docx(self, executor: DocumentToolExecutor, storage: AsyncMock):
        """Reading a .docx extracts paragraphs and tables."""
        storage.retrieve.return_value = _make_docx_bytes(["First paragraph", "Second"])
        result = await executor.execute("document_read", {"file_path": "/files/doc.docx"})

        assert "error" not in result
        assert result["format"] == "docx"
        assert "First paragraph" in result["paragraphs"]
        assert "Second" in result["paragraphs"]

    async def test_read_xlsx(self, executor: DocumentToolExecutor, storage: AsyncMock):
        """Reading an .xlsx returns sheet data."""
        storage.retrieve.return_value = _make_xlsx_bytes(
            [["Name", "Age"], ["Bob", 25]]
        )
        result = await executor.execute("document_read", {"file_path": "/files/data.xlsx"})

        assert "error" not in result
        assert result["format"] == "xlsx"
        assert "Sheet1" in result["sheets"]
        rows = result["sheets"]["Sheet1"]
        assert rows[0] == ["Name", "Age"]
        assert rows[1] == ["Bob", 25]

    async def test_read_pdf(self, executor: DocumentToolExecutor, storage: AsyncMock):
        """Reading a PDF extracts page text."""
        storage.retrieve.return_value = _make_pdf_bytes("Test PDF content")
        result = await executor.execute("document_read", {"file_path": "/files/report.pdf"})

        assert "error" not in result
        assert result["format"] == "pdf"
        assert result["page_count"] >= 1
        assert "Test PDF content" in result["text"]

    async def test_read_missing_path_returns_error(self, executor: DocumentToolExecutor):
        result = await executor.execute("document_read", {})
        assert "error" in result
        assert "file_path" in result["error"]

    async def test_read_unsupported_format_returns_error(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        storage.retrieve.return_value = b"data"
        result = await executor.execute("document_read", {"file_path": "/files/image.png"})
        assert "error" in result
        assert "Unsupported" in result["error"]


# ---------------------------------------------------------------------------
# Create tests
# ---------------------------------------------------------------------------


class TestDocumentCreate:
    async def test_create_docx(self, executor: DocumentToolExecutor, storage: AsyncMock):
        """Creating a .docx stores the file and returns the path."""
        result = await executor.execute(
            "document_create",
            {
                "format": "docx",
                "title": "Report",
                "content": {"paragraphs": ["Introduction", "Body"]},
            },
        )

        assert "error" not in result
        assert result["format"] == "docx"
        assert result["filename"] == "Report.docx"
        storage.store.assert_awaited_once()
        # Verify the stored bytes are a valid docx
        stored_bytes = storage.store.call_args[0][1]
        from docx import Document

        doc = Document(io.BytesIO(stored_bytes))
        texts = [p.text for p in doc.paragraphs]
        assert "Report" in texts  # heading
        assert "Introduction" in texts
        assert "Body" in texts

    async def test_create_xlsx(self, executor: DocumentToolExecutor, storage: AsyncMock):
        """Creating an .xlsx stores a workbook with the given sheets."""
        result = await executor.execute(
            "document_create",
            {
                "format": "xlsx",
                "title": "Data",
                "content": {
                    "sheets": [
                        {"name": "Sales", "rows": [["Q1", 100], ["Q2", 200]]},
                    ]
                },
            },
        )

        assert "error" not in result
        assert result["format"] == "xlsx"
        assert result["filename"] == "Data.xlsx"
        storage.store.assert_awaited_once()

        # Verify stored bytes are a valid xlsx
        from openpyxl import load_workbook

        stored_bytes = storage.store.call_args[0][1]
        wb = load_workbook(io.BytesIO(stored_bytes))
        assert "Sales" in wb.sheetnames
        ws = wb["Sales"]
        assert ws["A1"].value == "Q1"
        assert ws["B2"].value == 200
        wb.close()

    async def test_create_pdf(self, executor: DocumentToolExecutor, storage: AsyncMock):
        """Creating a PDF stores a valid PDF file."""
        result = await executor.execute(
            "document_create",
            {
                "format": "pdf",
                "title": "Summary",
                "content": {"paragraphs": ["This is the summary."]},
            },
        )

        assert "error" not in result
        assert result["format"] == "pdf"
        assert result["filename"] == "Summary.pdf"
        storage.store.assert_awaited_once()

        # Verify stored bytes start with PDF header
        stored_bytes = storage.store.call_args[0][1]
        assert stored_bytes[:5] == b"%PDF-"

    async def test_create_missing_format_returns_error(self, executor: DocumentToolExecutor):
        result = await executor.execute(
            "document_create", {"title": "X", "content": {}}
        )
        assert "error" in result

    async def test_create_missing_title_returns_error(self, executor: DocumentToolExecutor):
        result = await executor.execute(
            "document_create", {"format": "docx", "content": {}}
        )
        assert "error" in result

    async def test_create_unsupported_format_returns_error(
        self, executor: DocumentToolExecutor
    ):
        result = await executor.execute(
            "document_create",
            {"format": "bin", "title": "X", "content": {}},
        )
        assert "error" in result
        assert "Unsupported" in result["error"]


# ---------------------------------------------------------------------------
# Modify tests
# ---------------------------------------------------------------------------


class TestDocumentModify:
    async def test_modify_docx_append(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Appending paragraphs to an existing docx."""
        storage.retrieve.return_value = _make_docx_bytes(["Existing"])
        result = await executor.execute(
            "document_modify",
            {
                "file_path": "/files/doc.docx",
                "operation": "append",
                "data": {"paragraphs": ["New paragraph"]},
            },
        )

        assert "error" not in result
        assert result["operation"] == "append"
        assert result["status"] == "ok"
        storage.store.assert_awaited_once()

        # Verify the new document has both paragraphs
        from docx import Document

        stored_bytes = storage.store.call_args[0][1]
        doc = Document(io.BytesIO(stored_bytes))
        texts = [p.text for p in doc.paragraphs]
        assert "Existing" in texts
        assert "New paragraph" in texts

    async def test_modify_docx_replace(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Replacing text in a docx."""
        storage.retrieve.return_value = _make_docx_bytes(["Hello world, hello again"])
        result = await executor.execute(
            "document_modify",
            {
                "file_path": "/files/doc.docx",
                "operation": "replace",
                "data": {"old_text": "hello", "new_text": "goodbye"},
            },
        )

        assert "error" not in result
        assert result["operation"] == "replace"
        assert result["status"] == "ok"

    async def test_modify_xlsx_update_cells(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Updating cells in an xlsx."""
        storage.retrieve.return_value = _make_xlsx_bytes(
            [["Name", "Score"], ["Alice", 90]]
        )
        result = await executor.execute(
            "document_modify",
            {
                "file_path": "/files/data.xlsx",
                "operation": "update_cells",
                "data": {
                    "updates": [
                        {"sheet": "Sheet1", "cell": "B2", "value": 95},
                    ]
                },
            },
        )

        assert "error" not in result
        assert result["operation"] == "update_cells"
        assert result["updates_applied"] == 1
        assert result["status"] == "ok"

        # Verify the cell was updated
        from openpyxl import load_workbook

        stored_bytes = storage.store.call_args[0][1]
        wb = load_workbook(io.BytesIO(stored_bytes))
        assert wb["Sheet1"]["B2"].value == 95
        wb.close()

    async def test_modify_pdf_merge(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Merging multiple PDFs."""
        pdf1 = _make_pdf_bytes("Page one")
        pdf2 = _make_pdf_bytes("Page two")
        storage.retrieve.side_effect = [pdf1, pdf2]

        result = await executor.execute(
            "document_modify",
            {
                "file_path": "/files/main.pdf",
                "operation": "merge",
                "data": {"file_paths": ["/files/extra.pdf"]},
            },
        )

        assert "error" not in result
        assert result["operation"] == "merge"
        assert result["merged_count"] == 2
        assert result["status"] == "ok"

    async def test_modify_missing_path_returns_error(self, executor: DocumentToolExecutor):
        result = await executor.execute(
            "document_modify", {"operation": "append", "data": {}}
        )
        assert "error" in result

    async def test_modify_missing_operation_returns_error(
        self, executor: DocumentToolExecutor
    ):
        result = await executor.execute(
            "document_modify", {"file_path": "/files/doc.docx", "data": {}}
        )
        assert "error" in result

    async def test_modify_unsupported_operation_returns_error(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        storage.retrieve.return_value = _make_docx_bytes()
        result = await executor.execute(
            "document_modify",
            {
                "file_path": "/files/doc.docx",
                "operation": "delete_page",
                "data": {},
            },
        )
        assert "error" in result
        assert "Unsupported" in result["error"]

    async def test_modify_replace_missing_old_text(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Replace without old_text returns error."""
        storage.retrieve.return_value = _make_docx_bytes()
        result = await executor.execute(
            "document_modify",
            {
                "file_path": "/files/doc.docx",
                "operation": "replace",
                "data": {"new_text": "x"},
            },
        )
        assert "error" in result
        assert "old_text" in result["error"]


# ---------------------------------------------------------------------------
# Convert tests
# ---------------------------------------------------------------------------


class TestDocumentConvert:
    async def test_convert_xlsx_to_csv(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Converting xlsx to csv produces a valid CSV file."""
        storage.retrieve.return_value = _make_xlsx_bytes(
            [["Name", "Score"], ["Alice", 90], ["Bob", 85]]
        )
        result = await executor.execute(
            "document_convert",
            {"file_path": "/files/data.xlsx", "target_format": "csv"},
        )

        assert "error" not in result
        assert result["target_format"] == "csv"
        assert result["filename"] == "data.csv"
        storage.store.assert_awaited_once()

        # Verify the stored CSV content
        csv_bytes: bytes = storage.store.call_args[0][1]
        csv_text = csv_bytes.decode("utf-8")
        assert "Alice" in csv_text
        assert "Bob" in csv_text

    async def test_convert_docx_to_text(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Converting docx to text extracts plain text."""
        storage.retrieve.return_value = _make_docx_bytes(
            ["Line one", "Line two"]
        )
        result = await executor.execute(
            "document_convert",
            {"file_path": "/files/report.docx", "target_format": "text"},
        )

        assert "error" not in result
        assert result["target_format"] == "text"
        assert result["filename"] == "report.txt"
        storage.store.assert_awaited_once()

        text_bytes: bytes = storage.store.call_args[0][1]
        text = text_bytes.decode("utf-8")
        assert "Line one" in text
        assert "Line two" in text

    async def test_convert_missing_path_returns_error(
        self, executor: DocumentToolExecutor
    ):
        result = await executor.execute(
            "document_convert", {"target_format": "csv"}
        )
        assert "error" in result

    async def test_convert_missing_target_format_returns_error(
        self, executor: DocumentToolExecutor
    ):
        result = await executor.execute(
            "document_convert", {"file_path": "/files/data.xlsx"}
        )
        assert "error" in result

    async def test_convert_unsupported_conversion_returns_error(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        storage.retrieve.return_value = b"data"
        result = await executor.execute(
            "document_convert",
            {"file_path": "/files/image.png", "target_format": "csv"},
        )
        assert "error" in result
        assert "Unsupported" in result["error"]


# ---------------------------------------------------------------------------
# Executor dispatch tests
# ---------------------------------------------------------------------------


class TestExecutorDispatch:
    async def test_unknown_tool_returns_error(self, executor: DocumentToolExecutor):
        result = await executor.execute("document_unknown", {})
        assert "error" in result
        assert "Unknown" in result["error"]

    async def test_exception_in_handler_returns_error(
        self, executor: DocumentToolExecutor, storage: AsyncMock
    ):
        """Storage errors are caught and returned as error dicts."""
        storage.retrieve.side_effect = OSError("disk failure")
        result = await executor.execute(
            "document_read", {"file_path": "/files/broken.docx"}
        )
        assert "error" in result
        assert "disk failure" in result["error"]


# ---------------------------------------------------------------------------
# BuiltinToolExecutor delegation tests
# ---------------------------------------------------------------------------


class TestBuiltinDelegation:
    async def test_builtin_executor_delegates_document_tools(self):
        """BuiltinToolExecutor.execute delegates document_* calls to DocumentToolExecutor."""
        from flydesk.tools.builtin import BuiltinToolExecutor

        catalog_repo = MagicMock()
        catalog_repo.list_systems = AsyncMock(return_value=[])
        catalog_repo.list_endpoints = AsyncMock(return_value=[])
        audit_logger = MagicMock()
        audit_logger.query = AsyncMock(return_value=[])

        builtin = BuiltinToolExecutor(
            catalog_repo=catalog_repo, audit_logger=audit_logger
        )

        doc_executor = MagicMock()
        doc_executor.execute = AsyncMock(return_value={"status": "delegated"})
        builtin.set_document_executor(doc_executor)

        result = await builtin.execute("document_read", {"file_path": "/x.docx"})
        assert result == {"status": "delegated"}
        doc_executor.execute.assert_awaited_once_with(
            "document_read", {"file_path": "/x.docx"}
        )

    async def test_builtin_executor_without_doc_executor_returns_error(self):
        """Without a DocumentToolExecutor, document_* tools fall through to 'Unknown'."""
        from flydesk.tools.builtin import BuiltinToolExecutor

        catalog_repo = MagicMock()
        catalog_repo.list_systems = AsyncMock(return_value=[])
        catalog_repo.list_endpoints = AsyncMock(return_value=[])
        audit_logger = MagicMock()
        audit_logger.query = AsyncMock(return_value=[])

        builtin = BuiltinToolExecutor(
            catalog_repo=catalog_repo, audit_logger=audit_logger
        )
        # No doc executor set
        result = await builtin.execute("document_read", {"file_path": "/x.docx"})
        assert "error" in result
        assert "Unknown" in result["error"]
